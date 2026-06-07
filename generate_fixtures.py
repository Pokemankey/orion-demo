"""
Generate sample fixture files for local pipeline testing.
Creates fixtures/submission.json and three supporting documents
(PDF, DOCX, XLSX) covering all six risk dimensions.

Run with: python generate_fixtures.py
"""
import json
from pathlib import Path

from docx import Document as DocxDocument
from fpdf import FPDF
from openpyxl import Workbook


FIXTURES_DIR = Path("fixtures")
DOCS_DIR = FIXTURES_DIR / "docs"


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    _generate_annual_report_pdf()
    _generate_governance_policy_docx()
    _generate_risk_register_xlsx()
    _generate_submission_json()

    print("Fixtures generated:")
    for f in sorted(FIXTURES_DIR.rglob("*")):
        if f.is_file():
            print(f"  {f}")


# ---------------------------------------------------------------------------
# PDF — Annual Report (financial soundness, compliance, business model)
# ---------------------------------------------------------------------------

def _generate_annual_report_pdf() -> None:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Meridian Financial Services Ltd", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 10, "Annual Report 2024", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    _pdf_section(pdf, "Executive Summary",
        "Meridian Financial Services Ltd is a regulated payment processing and "
        "digital asset custody firm operating across the EU economic zone. In 2024 "
        "the firm processed EUR 2.4 billion in transaction volume across 18 jurisdictions. "
        "The firm employs 312 staff across offices in Frankfurt, Amsterdam, and Dublin."
    )

    _pdf_section(pdf, "Financial Position",
        "As of 31 December 2024, Meridian maintains a Tier 1 capital ratio of 8.2%, "
        "which falls below the recommended regulatory threshold of 10%. Management has "
        "acknowledged this shortfall and has committed to a capital raise of EUR 15 million "
        "in Q2 2025. The liquidity coverage ratio stands at 94%, against a minimum "
        "requirement of 100%. The leverage ratio is 4.1x, within acceptable bounds. "
        "Total assets: EUR 480 million. Total liabilities: EUR 395 million. "
        "Net profit for 2024: EUR 6.2 million, down 18% from EUR 7.6 million in 2023."
    )

    _pdf_section(pdf, "Regulatory and Compliance History",
        "In March 2023, ORION issued a formal warning to Meridian for a delay in "
        "suspicious transaction reporting under AML Directive Article 33. The firm "
        "remediated the reporting process within 60 days and no further action was taken. "
        "In September 2024, the firm received a minor administrative sanction of EUR 25,000 "
        "from the Dutch AFM for a late regulatory filing. No criminal proceedings or "
        "licence suspensions have occurred. The firm is currently subject to a routine "
        "supervisory review by the Frankfurt BaFin, expected to conclude in Q3 2025. "
        "KYC documentation coverage stands at 91% of active clients. The remaining 9% "
        "are flagged for enhanced due diligence review, which is ongoing."
    )

    _pdf_section(pdf, "Business Activities",
        "The firm operates across two primary business lines: (1) payment processing "
        "for e-commerce merchants, accounting for 68% of revenue, and (2) digital asset "
        "custody services for institutional clients, accounting for 32% of revenue. "
        "The top five clients account for 41% of total revenue, indicating moderate "
        "client concentration risk. The firm holds no proprietary trading positions. "
        "Declared activities for the current authorization period include: payment "
        "processing, digital asset custody, and foreign exchange conversion services."
    )

    pdf.add_page()
    _pdf_section(pdf, "Technology Infrastructure",
        "Core payment processing infrastructure is hosted on AWS EU-West. "
        "The firm completed a SOC 2 Type I audit in 2023 but has not yet completed "
        "the Type II audit, which was scheduled for late 2024 and is now delayed to "
        "mid-2025. No major security incidents were reported in 2024. "
        "The firm does not currently hold ISO 27001 certification."
    )

    out = DOCS_DIR / "annual_report.pdf"
    pdf.output(str(out))
    print(f"  Created {out}")


def _pdf_section(pdf: FPDF, title: str, body: str) -> None:
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, title, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=10)
    pdf.multi_cell(0, 6, body)
    pdf.ln(4)


# ---------------------------------------------------------------------------
# DOCX — Governance Policy (governance, operational resilience)
# ---------------------------------------------------------------------------

def _generate_governance_policy_docx() -> None:
    doc = DocxDocument()
    doc.add_heading("Meridian Financial Services Ltd", 0)
    doc.add_heading("Corporate Governance and Operational Policy", 1)

    doc.add_heading("Board Composition", 2)
    doc.add_paragraph(
        "The Board of Directors comprises five members: three executive directors "
        "(CEO, CFO, CRO) and two independent non-executive directors. "
        "The Chair is an independent non-executive director, Ms. A. Hoffman, appointed in 2021. "
        "The Board meets quarterly. An Audit Committee and a Risk Committee are in place, "
        "each chaired by an independent director. No director holds a material interest "
        "in any competitor firm. Beneficial ownership is held 51% by Meridian Group BV "
        "(Netherlands) and 49% by institutional investors. "
        "Ultimate beneficial owner: Mr. K. Brandt (German national), holding 34% through "
        "Meridian Group BV, declared in accordance with EU transparency requirements."
    )

    doc.add_heading("Remuneration Policy", 2)
    doc.add_paragraph(
        "Executive remuneration is subject to board approval and benchmarked annually "
        "against sector peers. Variable pay is capped at 100% of fixed salary in "
        "accordance with CRD V requirements. No guaranteed bonuses are offered. "
        "A clawback policy has been in place since 2022."
    )

    doc.add_heading("Business Continuity and Disaster Recovery", 2)
    doc.add_paragraph(
        "The firm maintains a Business Continuity Plan (BCP) reviewed annually, "
        "last reviewed January 2025. Recovery Time Objective (RTO) for core payment "
        "systems is 4 hours. Recovery Point Objective (RPO) is 1 hour. "
        "A full DR test was conducted in October 2024; all critical systems were "
        "restored within the defined RTO. The firm operates a hot standby environment "
        "in AWS EU-Central as a secondary site. Staff are trained in BCP procedures "
        "annually. The BCP covers pandemic, cyberattack, and infrastructure failure scenarios."
    )

    doc.add_heading("Incident Management", 2)
    doc.add_paragraph(
        "All operational incidents are logged in the firm's incident management system. "
        "In 2024, three P2 incidents were recorded, all resolved within SLA. "
        "No P1 (critical) incidents occurred. Average resolution time for P2 incidents: "
        "3.2 hours against an SLA of 6 hours. Incidents are reported to the Risk Committee "
        "on a quarterly basis. A post-incident review is mandatory for all P2 and above."
    )

    doc.add_heading("AML and KYC Framework", 2)
    doc.add_paragraph(
        "The firm operates a three-lines-of-defence model for AML compliance. "
        "Customer onboarding requires full KYC documentation including proof of identity, "
        "proof of address, and source of funds declaration for transactions above EUR 10,000. "
        "Enhanced Due Diligence is applied to all Politically Exposed Persons and "
        "high-risk jurisdictions as defined by FATF. Transaction monitoring is performed "
        "using automated rule-based screening. Suspicious Activity Reports are filed "
        "with the relevant Financial Intelligence Unit within 24 hours of identification. "
        "The MLRO position is held by Ms. C. Dijkstra, appointed 2022."
    )

    out = DOCS_DIR / "governance_policy.docx"
    doc.save(str(out))
    print(f"  Created {out}")


# ---------------------------------------------------------------------------
# XLSX — Risk Register (operational resilience, cybersecurity)
# ---------------------------------------------------------------------------

def _generate_risk_register_xlsx() -> None:
    wb = Workbook()

    # Sheet 1 — Operational Risk Register
    ws1 = wb.active
    ws1.title = "Operational Risk Register"
    ws1.append(["Risk ID", "Category", "Description", "Likelihood", "Impact", "Mitigation", "Status"])
    ws1.append(["OP-001", "Operational", "Key person dependency on CTO role", "Medium", "High",
                "Succession plan documented; deputy CTO nominated", "Open"])
    ws1.append(["OP-002", "Operational", "Third-party payment processor outage", "Low", "High",
                "Dual processor agreement in place with fallback to backup processor", "Mitigated"])
    ws1.append(["OP-003", "Technology", "Cloud provider regional outage (AWS EU-West)", "Low", "Critical",
                "Hot standby in EU-Central; RTO 4 hours tested October 2024", "Mitigated"])
    ws1.append(["OP-004", "Compliance", "Late regulatory filing", "Low", "Medium",
                "Automated filing calendar with 14-day advance alerts implemented Q1 2025", "Mitigated"])
    ws1.append(["OP-005", "Cybersecurity", "Phishing attack on staff credentials", "High", "High",
                "MFA enforced; phishing simulation training quarterly", "Open"])
    ws1.append(["OP-006", "Cybersecurity", "Ransomware on internal systems", "Medium", "Critical",
                "Endpoint detection deployed; offline backups maintained", "Open"])

    # Sheet 2 — Key Risk Indicators
    ws2 = wb.create_sheet("Key Risk Indicators")
    ws2.append(["KRI", "Threshold", "Current Value", "Status", "Period"])
    ws2.append(["Transaction error rate", "< 0.05%", "0.03%", "Green", "Q4 2024"])
    ws2.append(["System availability", "> 99.9%", "99.94%", "Green", "Q4 2024"])
    ws2.append(["SAR filing timeliness", "100% within 24h", "98.2%", "Amber", "Q4 2024"])
    ws2.append(["KYC completion rate", "> 95%", "91%", "Red", "Q4 2024"])
    ws2.append(["Staff AML training completion", "> 95%", "97%", "Green", "Q4 2024"])
    ws2.append(["Open high-severity audit findings", "< 3", "1", "Green", "Q4 2024"])

    out = DOCS_DIR / "risk_register.xlsx"
    wb.save(str(out))
    print(f"  Created {out}")


# ---------------------------------------------------------------------------
# Submission JSON
# ---------------------------------------------------------------------------

def _generate_submission_json() -> None:
    submission = {
        "applicant_id": "FIRM-2024-0042",
        "firm_name": "Meridian Financial Services Ltd",
        "declared_activities": [
            "payment_processing",
            "digital_asset_custody",
            "foreign_exchange_conversion"
        ],
        "documents": [
            # .as_posix() keeps forward slashes on every platform — these are
            # storage URIs, not OS paths, and must resolve in the Linux container.
            {"type": "pdf",  "path": (DOCS_DIR / "annual_report.pdf").as_posix()},
            {"type": "docx", "path": (DOCS_DIR / "governance_policy.docx").as_posix()},
            {"type": "xlsx", "path": (DOCS_DIR / "risk_register.xlsx").as_posix()},
        ]
    }

    out = FIXTURES_DIR / "submission.json"
    out.write_text(json.dumps(submission, indent=2), encoding="utf-8")
    print(f"  Created {out}")


if __name__ == "__main__":
    main()
